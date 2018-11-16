"""
This module integrates the system with other systems and 
parses data to the correct output format.
"""
# No integration with News BD yet

import copy  # Table extraction
import datetime
import io  # image from url
import os # multiplatform integration
from pathlib import Path  # multiplatform integration
import requests  # Picture recover from url

from docx import Document  # Accessing and creating documents
from docx.enum.style import WD_STYLE_TYPE  # style types
from docx.enum.text import WD_ALIGN_PARAGRAPH  # paragraph alignment
from docx.shared import Cm, Pt, RGBColor  # Image and font sizes

# Docx Output

# template document name
ORIGIN = os.path.join('interface', 'src', 'assets', 'template.docx')

# colors RGB values
DARK_GRAY = [0x68, 0x68, 0x68]
BLACK = [0, 0, 0]

def img_from_url(url):
    """
    Returns a binary image from a url
    """
    response = requests.get(url)  
    # Create the picture
    binary_img = io.BytesIO(response.content)  
    return binary_img

def get_template_table(origin):
    """
    Gives a copy of the template table in origin file
    """
    template = Document(origin)
    # Extract template table
    table_skeleton = template.tables[0]
    return copy.deepcopy(table_skeleton)

def add_table(output, table):
    """
    Adds a table to output file.
    """
    table = table._tbl # No idea. Just works
    paragraph = output.add_paragraph() # Create a new paragraph to hold the table
    paragraph._p.addnext(table) # After that, we add the previously copied table

def add_style(doc, style_name, font_family, font_size, rgb, style_type=WD_STYLE_TYPE.CHARACTER, bold=False):
    """
    Creates a new style to be applyed to objects.
    Returns style name.
    """
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style(style_name, style_type)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(font_size)
    obj_font.name = font_family
    obj_font.bold = bold
    obj_font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
    return style_name

def change_margins(doc, margin):
    """
    Changes margins of a document to specific size.
    """
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)

def format_table(table, article):
    """
    Populates the table correctly from article.
    """
    region = table.cell(0, 2).paragraphs[0]
    title = table.cell(1, 1).paragraphs[0]
    img = table.cell(1, 3).paragraphs[0]
    description = table.cell(2, 1).paragraphs[0]
    thumbnail = table.cell(2, 3).paragraphs[0]
    url = table.cell(3, 0).paragraphs[0]

    region.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Security checks
    art_reg = (article.region if isinstance(article.region, str) else "Região da Notícia")
    art_title = (article.title if isinstance(article.title, str) else "Título da Notícia")
    art_description = (article.description if isinstance(article.description, str) else "Descrição da Notícia")
    art_url = (article.title if isinstance(article.url, str) else "URL da Noticia")

    # Text info
    region.add_run(art_reg, style='regionStyle').bold = True 
    title.add_run(art_title, style='titleStyle').bold = True 
    description.add_run(art_description, style='textStyle')
    url.add_run(art_url, style='urlStyle')

    # Image info
    if isinstance(article.url_to_image, str):
        pic = img_from_url(article.url_to_image)
        img.add_run().add_picture(pic, width=Cm(6), height=Cm(5))

    # thumbnail not working yet
    thumbnail.add_run('tiny img')

def add_centered_img(doc, path, width):
    """
    Adds a new picture to doc, centered.
    """
    doc.add_picture(path, width=width)
    footer_img = doc.paragraphs[-1] # last paragraph, aka the picture
    footer_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_footer(doc):
    """
    Adds the clipping's footer.
    """
    path_to_img = os.path.join('interface', 'src', 'assets', 'sala_logo.png')
    doc.add_picture(path_to_img, width=Cm(2))
    footer_img = doc.paragraphs[-1] # last paragraph, aka the picture
    footer_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    people_involved = ['Elaboração', 'Tradução', 'Equipe Editorial', 'Revisão']
    for person in people_involved:
        par = doc.add_paragraph(person, style='footerStyle')
        par.bold = True
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_docx(articles, path_to_output):
    """
    Creates a docx file with tables for each of the article in articles.
    """
    output_name = 'Clipping_' + datetime.date.today().strftime('%a_%d%b%Y')
    output = Document()

    # Adds region bars centered
    path = os.path.join('interface', 'src', 'assets')
    add_centered_img(output, os.path.join(path, 'brasil_bar.png'), (output.sections[0].page_width * 0.8))
    add_centered_img(output, os.path.join(path, 'mundo_bar.png'), (output.sections[0].page_width * 0.8))

    # Create character styles
    add_style(output, 'regionStyle', 'PT Sans', 14, DARK_GRAY)
    add_style(output, 'titleStyle', 'PT Sans', 14, BLACK)
    add_style(output, 'textStyle', 'PT Sans', 12, BLACK)
    add_style(output, 'urlStyle', 'PT Sans', 8, DARK_GRAY)
    # Paragraph style
    add_style(output, 'footerStyle', 'PT Sans', 10, BLACK, WD_STYLE_TYPE.PARAGRAPH, True)

    # Expand margins
    change_margins(output, 1)

    # Copy tables template to output file
    for _ in range(len(articles)):
        add_table(output, get_template_table(ORIGIN))
    # Populate tables
    for idx, article in enumerate(articles):
        print(idx)
        table = output.tables[idx]
        format_table(table, article)

    # Footer
    add_footer(output)

    # Save doc
    path = Path(path_to_output) / (output_name + '.docx')
    output.save(path)
    return path
