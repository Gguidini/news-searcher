"""
This module integrates the system with other systems and
parses data to the correct output format.

Selected News from search are pushed to News DB,
and formatted into a docx document for further editing by hand.
Extensive use of python-docx library. You might want to get familiarized with it.
"""

import copy  # Table extraction
import datetime # docx name
import io  # image from url
import os # multiplatform integration - paths
import requests  # Picture recover from url

from bs4 import BeautifulSoup # Extract icon url
from docx import Document  # Accessing and creating documents
from docx.enum.style import WD_STYLE_TYPE  # style types
from docx.enum.text import WD_ALIGN_PARAGRAPH  # paragraph alignment
from docx.shared import Cm, Pt, RGBColor  # Image and font sizes

from interface.src.config import BD_INFO
BD_URL, BD_PASSWD = BD_INFO()

# News DB Output ==============================
def push_to_DB(article):
    """
    Transforms a News article into a json object,
    inserts object in the News data base.
    ps.: database is not part of this system.
    """
    response = requests.get(BD_URL, params={'json':article.to_json(), 'key':BD_PASSWD})
    return response

# Docx Output =================================

# template document path
ORIGIN = os.path.join('interface', 'src', 'assets', 'template.docx')

# base document path
BASE = os.path.join('interface', 'src', 'assets', 'base.docx')

# colors RGB values
DARK_GRAY = [0x68, 0x68, 0x68]
BLACK = [0, 0, 0]

def _img_from_url_(url: str):
    """
    Returns a binary image from a url.
    
    argument:
    > url - string
        url to News image. News.url_to_image.
    """
    response = requests.get(url)  
    # Create the picture
    binary_img = io.BytesIO(response.content)  
    return binary_img

def _icon_from_url_(url: str):
    """
    Returns the url to the icon of a webpage.
    Icon will later be added to the docx table.
    """
    page = requests.get(url)
    soup = BeautifulSoup(page.text, features='lxml')
    for string in ['icon', 'Icon', 'ICON', 'shortcut icon']: # soup is case sensitive
        icon_link = soup.find("link", rel=string)
        if icon_link is not None:
            return icon_link['href']
    return None

def _get_template_table_(origin):
    """
    Gives a true copy of the template table in ORIGIN file.
    Template table created by Sala de Situação's team.

    argument:
    > origin - string
        path to ORIGIN file.
    """
    template = Document(origin)
    # Extract template table
    table_skeleton = template.tables[0]
    return copy.deepcopy(table_skeleton)

def _add_table_(output, table):
    """
    Adds a table to output file.

    arguments:
    > output : Document()
        A docx document to put the table into.
    
    > table : Table()
        A Table object that will be inserted in output.
    """
    table = table._tbl # No idea. Just works
    paragraph = output.add_paragraph() # Create a new paragraph to hold the table
    paragraph._p.addnext(table) # After that, we add the previously copied table

def _add_style_(doc, style_name, font_family, font_size, rgb, style_type=WD_STYLE_TYPE.CHARACTER, bold=False):
    """
    Creates a new style to be applyed to objects.
    Returns style name.

    arguments:
    > doc : Document
        Styles are bound to Documents. The Document to add the style.
    > style_name : str
        Any name that will reference this style later.
    > font_family : str
        The font family. Must be a font that exists in Word.
    > font_size : int
        The font size
    > rgb : Tuple[int, int, int]
        Font color, in RGB format.
    > style_type
        Whether the style is to be applied to runs or paragraphs. Default is run.
        WD_STYLE_TYPE.CHARACTER -> runs
        WD_STYLE_TYPE.PARAGRAPH -> paragraphs
    > bold : bool
        Whether style is bolded or not.
    """
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style(style_name, style_type)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(font_size)
    obj_font.name = font_family
    obj_font.bold = bold
    obj_font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
    return style_name

def _change_margins_(doc, margin):
    """
    Changes margins of a document to specific size.

    arguments:
    > doc : Document
        Document to modify
    > margin : int
        Size of margin, in centimeters. All margin will have the same size.
    """
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)

def _format_table_(table, article):
    """
    Populates the table correctly from article.
    Follows table template from ORIGIN document.
    """
    region = table.cell(0, 2).paragraphs[0]
    title = table.cell(1, 1).paragraphs[0]
    img = table.cell(1, 3).paragraphs[0]
    description = table.cell(2, 1).paragraphs[0]
    thumbnail = table.cell(2, 3).paragraphs[0]
    url = table.cell(3, 0).paragraphs[0]

    region.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Security checks
    art_reg = (article.region if isinstance(article.region, str) else "Região da Notícia")
    art_title = (article.title if isinstance(article.title, str) else "Título da Notícia")
    art_description = (article.description if isinstance(article.description, str) else "Descrição da Notícia")
    art_url = (article.url if isinstance(article.url, str) else "URL da Noticia")

    # Text info
    region.add_run(art_reg, style='regionStyle').bold = True 
    title.add_run(art_title, style='titleStyle').bold = True 
    description.add_run(art_description, style='textStyle')
    url.add_run(art_url, style='urlStyle')

    # Image info
    try:
        pic = _img_from_url_(article.url_to_image)
        img.add_run().add_picture(pic, width=Cm(6), height=Cm(5))
    except:
        img.add_run("Image not found")

    # thumbnail is the shortcut icon of the News URL, if any.
    thumbnail_url = _icon_from_url_(article.url)
    try:
        thumbnail_pic = _img_from_url_(thumbnail_url)
        thumbnail.add_run().add_picture(thumbnail_pic, width=Cm(2.5), height=Cm(2.5))
    except:
        thumbnail.add_run('Image not found')

def _add_centered_img_(doc, path, width):
    """
    Adds a new picture to doc, centered.
    Picture must exist in the computer.

    arguments:
    > doc : Document
        Document to modify
    > path
        Path to image
    """
    doc.add_picture(path, width=width)
    footer_img = doc.paragraphs[-1] # last paragraph, aka the picture
    footer_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _add_footer_(doc):
    """
    Adds the clipping's footer.
    Copied from clipping template provided by Sala de Situação.
    """
    path_to_img = os.path.join('interface', 'src', 'assets', 'sala_logo.png')
    doc.add_picture(path_to_img, width=Cm(2))
    footer_img = doc.paragraphs[-1] # last paragraph, aka the picture
    footer_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    people_involved = ['Elaboração', 'Tradução', 'Equipe Editorial', 'Revisão']
    for person in people_involved:
        par = doc.add_paragraph(person, style='footerStyle')
        par.bold = True
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_docx(articles, path_to_output):
    """
    Creates a docx file with tables for each of the article in articles.
    Follows template provided by Sala de Situação.
    """
    output_name = 'Clipping_' + datetime.date.today().strftime('%a_%d%b%Y')
    output = Document(BASE)

    # Adds region bars centered
    path = os.path.join('interface', 'src', 'assets')
    _add_centered_img_(output, os.path.join(path, 'brasil_bar.png'),
                     (output.sections[0].page_width * 0.8))
    _add_centered_img_(output, os.path.join(path, 'mundo_bar.png'), 
                     (output.sections[0].page_width * 0.8))

    # Create character styles
    _add_style_(output, 'regionStyle', 'PT Sans', 14, DARK_GRAY)
    _add_style_(output, 'titleStyle', 'PT Sans', 14, BLACK)
    _add_style_(output, 'textStyle', 'PT Sans', 12, BLACK)
    _add_style_(output, 'urlStyle', 'PT Sans', 8, DARK_GRAY)
    # Paragraph style
    _add_style_(output, 'footerStyle', 'PT Sans', 10, BLACK, WD_STYLE_TYPE.PARAGRAPH, True)

    # Expand margins
    _change_margins_(output, 1)

    # Copy tables template to output file
    for _ in range(len(articles)):
        _add_table_(output, _get_template_table_(ORIGIN))
    # Populate tables
    for idx, article in enumerate(articles):
        table = output.tables[idx]
        _format_table_(table, article)

    # Footer
    _add_footer_(output)

    # Save doc
    path = os.path.join(path_to_output,  output_name + '.docx')
    output.save(path)
    return output_name + '.docx'
